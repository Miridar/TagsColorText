import os
import re
from docx import Document
from docx.shared import RGBColor

# Color mapping
color_mapping = {
    "adjective": RGBColor(255, 192, 203),  # Pink
    "adverb": RGBColor(255, 255, 0),       # Yellow
    "conjunction": RGBColor(192, 192, 192),# Grey
    "determiner": RGBColor(255, 0, 255),   # Magenta
    "noun": RGBColor(0, 255, 0),           # Green
    "number": RGBColor(255, 165, 0),       # Orange
    "preposition": RGBColor(255, 105, 180),# Hot pink
    "pronoun": RGBColor(175, 238, 238),    # Pale turquoise
    "verb": RGBColor(0, 0, 255)            # Blue
}

# Determine the directory of the script
script_dir = os.path.dirname(os.path.abspath(__file__))
input_file_path = os.path.join(script_dir, "input.txt")

# Load the text from "input.txt"
with open(input_file_path, "r", encoding="utf-8") as file:
    text_with_tags = file.read()

# Function to create a word document with colored words
def create_colored_doc(text, colors):
    # Determine the directory of the script
    script_dir = os.path.dirname(os.path.abspath(__file__))

    # Create a new Document
    doc = Document()

    # Split the text into parts (tagged and untagged)
    parts = re.split(r'(<\w+>[^<]+<\/\w+>)', text)
    
    p = doc.add_paragraph()

    for part in parts:
        if part != "":
            # Check if the part is a tagged word
            match = re.match(r'<(\w+)>([^<]+)<\/\w+>', part)
            if match:
                tag, word = match.groups()
                run = p.add_run(word)  # Add a space after each word for separation
                run.font.color.rgb = colors.get(tag, RGBColor(0, 0, 0))
            else:
                # Part of the text without tags
                run = p.add_run(part + " ")
                run.font.color.rgb = RGBColor(0, 0, 0)  # Default color for untagged text

    # Construct the path for the Word document in the script's directory
    doc_path = os.path.join(script_dir, "colored_text.docx")
    doc.save(doc_path)

    return doc_path

# Create the HTML document with colored text
html_path = create_colored_doc(text_with_tags, color_mapping)
